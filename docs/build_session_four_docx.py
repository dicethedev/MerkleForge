#!/usr/bin/env python3.12
"""Build a self-contained DOCX for Session Four."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets" / "session-four"
MARKDOWN = DOCS / "session-four-implementation-and-testing.md"
OUT = MARKDOWN.with_suffix(".docx")

CODE_SHOTS = [
    ("Cargo.toml", ROOT / "Cargo.toml", 1, 36, "code-workspace-manifest.png"),
    (
        "merkle_tree.rs",
        ROOT / "merkle-core" / "src" / "traits" / "merkle_tree.rs",
        15,
        76,
        "code-merkle-tree-trait.png",
    ),
    (
        "sparse_properties.rs",
        ROOT / "merkle-variants" / "tests" / "sparse_properties.rs",
        58,
        111,
        "code-sparse-properties.png",
    ),
    ("ci.yml", ROOT / ".github" / "workflows" / "ci.yml", 1, 85, "code-ci-workflow.png"),
]


def load_font(size: int) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSansMono.ttf",
        "/usr/share/fonts/truetype/liberation2/LiberationMono-Regular.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def render_code_shot(
    title: str,
    source: Path,
    start: int,
    end: int,
    output_name: str,
) -> None:
    lines = source.read_text().splitlines()
    selected = lines[start - 1 : end]
    numbered = [f"{line_no:>3}  {line}" for line_no, line in enumerate(selected, start)]

    font = load_font(22)
    title_font = load_font(24)
    line_height = 32
    padding_x = 34
    padding_y = 30
    title_gap = 48
    max_chars = max(len(line) for line in numbered)
    width = min(1800, max(980, padding_x * 2 + max_chars * 13))
    height = padding_y * 2 + title_gap + len(numbered) * line_height

    image = Image.new("RGB", (width, height), "#0b1017")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle(
        (8, 8, width - 8, height - 8),
        radius=26,
        fill="#101722",
        outline="#2e74b5",
        width=2,
    )
    draw.text((padding_x, padding_y), title, fill="#dbeafe", font=title_font)
    draw.line(
        (padding_x, padding_y + 34, width - padding_x, padding_y + 34),
        fill="#26364a",
        width=2,
    )

    y = padding_y + title_gap
    for line in numbered:
        if line.strip().startswith("//") or line.strip().startswith("#"):
            fill = "#8aa4bf"
        elif any(keyword in line for keyword in ("pub ", "fn ", "trait ", "impl ", "use ")):
            fill = "#b9f6ff"
        else:
            fill = "#e5edf7"
        draw.text((padding_x, y), line[:132], fill=fill, font=font)
        y += line_height

    image.save(ASSETS / output_name)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text.strip())
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(8.5)
    if bold:
        run.font.color.rgb = RGBColor(11, 37, 69)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_widths(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(width)


def add_table(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = len(rows[0])
    table = document.add_table(rows=1, cols=col_count)
    table.style = "Table Grid"

    for idx, text in enumerate(rows[0]):
        set_cell_shading(table.rows[0].cells[idx], "F2F4F7")
        set_cell_text(table.rows[0].cells[idx], text, bold=True)

    for row_values in rows[1:]:
        cells = table.add_row().cells
        for idx, text in enumerate(row_values):
            set_cell_text(cells[idx], text)

    if col_count == 2:
        widths = [2.0, 4.2]
    elif col_count == 3:
        widths = [1.7, 2.25, 2.25]
    elif col_count == 4:
        widths = [1.35, 1.6, 1.85, 1.6]
    elif col_count == 5:
        widths = [1.0, 1.55, 1.65, 1.55, 1.0]
    else:
        widths = [6.2 / col_count] * col_count
    set_table_widths(table, widths)

    document.add_paragraph()


def add_code_block(document: Document, code: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.left_indent = Inches(0.12)
    for line_idx, line in enumerate(code.rstrip().splitlines()):
        if line_idx:
            paragraph.add_run("\n")
        run = paragraph.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(31, 41, 55)


def add_image(document: Document, alt: str, rel_path: str) -> None:
    image_path = MARKDOWN.parent / rel_path
    if not image_path.exists():
        return
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    try:
        run.add_picture(str(image_path), width=Inches(6.2))
    except Exception:
        run.add_text(f"[Image omitted: {rel_path}]")

    caption = document.add_paragraph(alt)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    caption.runs[0].italic = True
    caption.runs[0].font.size = Pt(9)
    caption.runs[0].font.color.rgb = RGBColor(85, 85, 85)


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        line = lines[idx].strip()
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            idx += 1
            continue
        rows.append(cells)
        idx += 1
    return rows, idx


def add_markdown(document: Document, text: str) -> None:
    lines = text.splitlines()
    idx = 0
    in_code = False
    code_lines: list[str] = []
    list_buffer: list[tuple[str, str]] = []

    def flush_list() -> None:
        nonlocal list_buffer
        for kind, item in list_buffer:
            if kind == "number":
                paragraph = document.add_paragraph(item)
                paragraph.paragraph_format.left_indent = Inches(0.25)
                paragraph.paragraph_format.first_line_indent = Inches(-0.18)
            else:
                document.add_paragraph(item, style="List Bullet")
        list_buffer = []

    while idx < len(lines):
        raw = lines[idx]
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(document, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                flush_list()
                in_code = True
            idx += 1
            continue

        if in_code:
            code_lines.append(raw)
            idx += 1
            continue

        if not line.strip() or line.strip() == "---":
            flush_list()
            idx += 1
            continue

        image = re.fullmatch(r"!\[(.*?)\]\((.*?)\)", line.strip())
        if image:
            flush_list()
            add_image(document, image.group(1), image.group(2))
            idx += 1
            continue

        if line.strip().startswith("|") and idx + 1 < len(lines) and lines[idx + 1].strip().startswith("|"):
            flush_list()
            table_rows, idx = parse_table(lines, idx)
            add_table(document, table_rows)
            continue

        heading = re.match(r"^(#{1,4})\s+(.*)$", line)
        if heading:
            flush_list()
            level = len(heading.group(1))
            text_value = strip_inline_markdown(heading.group(2))
            if level == 1:
                document.add_heading(text_value, level=0)
            else:
                document.add_heading(text_value, level=min(level - 1, 3))
            idx += 1
            continue

        bullet = re.match(r"^\s*-\s+(.*)$", line)
        number = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
        if bullet:
            list_buffer.append(("bullet", strip_inline_markdown(bullet.group(1))))
            idx += 1
            continue
        if number:
            list_buffer.append(("bullet", strip_inline_markdown(number.group(2))))
            idx += 1
            continue

        flush_list()
        paragraph_lines = [line.strip()]
        idx += 1
        while idx < len(lines):
            nxt = lines[idx].rstrip()
            if (
                not nxt.strip()
                or nxt.startswith("#")
                or nxt.startswith("```")
                or nxt.strip().startswith("|")
                or nxt.strip().startswith("- ")
                or re.match(r"^\s*\d+\.\s+", nxt)
                or re.fullmatch(r"!\[(.*?)\]\((.*?)\)", nxt.strip())
            ):
                break
            paragraph_lines.append(nxt.strip())
            idx += 1
        add_rich_paragraph(document, " ".join(paragraph_lines))


def strip_inline_markdown(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    text = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", text)
    return text


def add_rich_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(7)
    parts = re.split(r"(\*\*.*?\*\*|`[^`]*`|\[.*?\]\(.*?\))", text)
    for part in parts:
        if not part:
            continue
        run_text = strip_inline_markdown(part)
        run = paragraph.add_run(run_text)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        if part.startswith("**") and part.endswith("**"):
            run.bold = True
        if part.startswith("`") and part.endswith("`"):
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(11, 37, 69)


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Title", 24, RGBColor(11, 37, 69)),
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = style_name != "Title"


def add_front_matter(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Session Four\nImplementation and Testing")
    run.bold = True
    run.font.name = "Calibri"
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("MerkleForge Project Report")
    run.font.name = "Calibri"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(85, 85, 85)

    summary = document.add_paragraph()
    summary.alignment = WD_ALIGN_PARAGRAPH.CENTER
    summary.paragraph_format.space_after = Pt(24)
    run = summary.add_run(
        "Prepared as implementation, testing, performance, and user-operation evidence for the MerkleForge Rust workspace."
    )
    run.font.name = "Calibri"
    run.font.size = Pt(11)

    document.add_page_break()


def add_footer(document: Document) -> None:
    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("MerkleForge - Session Four Implementation and Testing")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(85, 85, 85)


def main() -> None:
    ASSETS.mkdir(parents=True, exist_ok=True)
    for shot in CODE_SHOTS:
        render_code_shot(*shot)

    markdown_text = MARKDOWN.read_text()
    marker = "Recommended codebase screenshots to include in the final report:"
    insertion = """
### Codebase Screenshots

The following screenshots show selected parts of the implementation and CI
configuration used during system development.

![Figure 4.8: Workspace manifest and shared dependency versions](assets/session-four/code-workspace-manifest.png)

![Figure 4.9: Shared MerkleTree and ProofVerifier trait contract](assets/session-four/code-merkle-tree-trait.png)

![Figure 4.10: Sparse tree property-based tests](assets/session-four/code-sparse-properties.png)

![Figure 4.11: GitHub Actions CI workflow](assets/session-four/code-ci-workflow.png)

"""
    markdown_text = markdown_text.replace(marker, insertion + marker)

    document = Document()
    configure_styles(document)
    add_front_matter(document)
    add_markdown(document, markdown_text)
    add_footer(document)
    document.save(OUT)


if __name__ == "__main__":
    main()
